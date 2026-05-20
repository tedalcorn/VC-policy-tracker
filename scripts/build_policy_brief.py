#!/usr/bin/env python3
"""Generate a shareable policy-asks brief as a .docx for Ted to send to editors.

Filters articles with AI-assigned policy-ask score >= MIN_SCORE, groups by the
article's primary domain (first bucket), and within Public Safety sub-groups
by who is being asked. Within each (sub)group: bullets with date, title (as
hyperlink), author, and the ask in italics with the bolded target prefix.

Output: projects/2026-05-19-policy-asks-brief/brief.docx
"""

from __future__ import annotations

import csv
import json
import re
import sys
from collections import defaultdict
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt, RGBColor

ROOT = Path(__file__).resolve().parent.parent
SCORES = ROOT / "data" / "processed" / "policy_scores.csv"
ARTICLES = ROOT / "data" / "processed" / "articles.json"
OUT = ROOT / "projects" / "2026-05-19-policy-asks-brief" / "brief.docx"

MIN_SCORE = 4

DOMAIN_ORDER = [
    "Public Safety",
    "Governance",
    "Politics",
    "Housing",
    "Transit",
    "Health & human services",
    "Youth",
    "Economy",
    "Culture",
    "Data",
    "Race",
    "Immigration",
    "Climate",
    "Other",
]

# Ordered most-specific first; first match wins.
PUBLIC_SAFETY_SUBGROUPS = [
    ("Department of Community Safety (Mamdani administration)",
        re.compile(r"\bDepartment of Community Safety\b|\bcommunity safety office\b", re.I)),
    ("Rikers, jails, and the courts that govern them",
        re.compile(r"\bRikers\b|\bjails?\b|\bDOC\b|\bDepartment of Correction\b|\breceiver\b|\bNunez\b|\bSwain\b|\bcorrections?\b", re.I)),
    ("NYPD",
        re.compile(r"\bNYPD\b|\bpolice department\b|\bpolice strategy\b", re.I)),
    ("State government (Albany / Hochul / Legislature / state agencies)",
        re.compile(r"\bAlbany\b|\bHochul\b|\bstate legislature\b|\bstate assembly\b|\bstate senate\b|\bstate government\b|\bstate court system\b|\bstate of new york\b|\bgovernor\b", re.I)),
    ("Adams administration (may no longer be relevant)",
        re.compile(r"\bMayor Adams\b|\bAdams administration\b|\bEric Adams\b", re.I)),
    ("Other city government (Mayor's office, City Council, etc.)",
        re.compile(r"\bCity Council\b|\bMayor\b|\bcity hall\b|\bcity government\b|\bnext mayoral administration\b|\bnew york city government\b", re.I)),
]


def fmt_date(iso: str) -> str:
    if not iso or len(iso) < 10:
        return iso
    y, m, d = iso[:4], iso[5:7], iso[8:10]
    return f"{int(m)}/{int(d)}/{y[2:]}"


def freshen_mamdani(s: str) -> str:
    """Mamdani is mayor as of 2026, so promote any older Mayor-elect / next-mayor
    phrasings now that the work is being read in the present."""
    if not s:
        return s
    out = s
    out = re.sub(r"\bMayor[- ]elect (Zohran )?Mamdani\b", "Mayor Mamdani", out)
    out = re.sub(r"\bZohran Mamdani\b(?!')", "Mayor Mamdani", out)
    return out


def classify_public_safety(target: str, ask: str) -> str:
    haystack = (target or "") + " " + (ask or "")
    for label, pat in PUBLIC_SAFETY_SUBGROUPS:
        if pat.search(haystack):
            return label
    return "Other (federal, regulatory, etc.)"


def add_hyperlink(paragraph, url: str, text: str):
    """Insert a clickable hyperlink into a paragraph (python-docx has no
    direct API for this, so we wire the relationship + XML by hand)."""
    part = paragraph.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    new_run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    color = OxmlElement("w:color"); color.set(qn("w:val"), "0563C1"); rPr.append(color)
    underline = OxmlElement("w:u"); underline.set(qn("w:val"), "single"); rPr.append(underline)
    new_run.append(rPr)
    t = OxmlElement("w:t")
    t.text = text
    t.set(qn("xml:space"), "preserve")
    new_run.append(t)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


def add_ask_paragraph(doc, q: dict) -> None:
    """A single-line bullet per article. Date and title-link on top; the
    ask italicized below. Target isn't repeated separately — the ask itself
    already names it as the subject in most cases."""
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.1
    p.add_run(f"{fmt_date(q['date'])} · ").bold = True
    add_hyperlink(p, q["url"], q["title"])
    p.add_run(f" — {q['author']}")

    cont = doc.add_paragraph()
    cont.paragraph_format.left_indent = Pt(18)
    cont.paragraph_format.space_before = Pt(0)
    cont.paragraph_format.space_after = Pt(4)
    cont.paragraph_format.line_spacing = 1.1
    cont.add_run(q["ask"]).italic = True


def main() -> int:
    by_slug = {a["slug"]: a for a in json.loads(ARTICLES.read_text())}
    with SCORES.open(newline="") as f:
        scored = list(csv.DictReader(f))

    qualifying = []
    for r in scored:
        try:
            s = int(r["score"])
        except (ValueError, TypeError):
            continue
        if s < MIN_SCORE:
            continue
        art = by_slug.get(r["slug"])
        if not art or art.get("article_type") != "article":
            continue
        primary = (art.get("buckets") or ["Other"])[0]
        qualifying.append({
            "score": s,
            "slug": r["slug"],
            "title": r["title"],
            "date": r["published_at"][:10],
            "author": r["primary_author"],
            "url": art.get("url", ""),
            "target": freshen_mamdani(r["target"] or "(no specific target named)"),
            "ask": freshen_mamdani(r["policy_ask"]),
            "domain": primary,
        })

    by_domain: dict[str, list[dict]] = defaultdict(list)
    for q in qualifying:
        by_domain[q["domain"]].append(q)

    total = len(qualifying)

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    doc.add_heading("Vital City Policy Asks — Where We've Been Calling for Action", level=0)

    intro = doc.add_paragraph()
    r = intro.add_run(
        f"Generated from an AI-assisted scan of every Vital City article "
        f"published between Sept 2021 and May 2026. {total} of 663 articles "
        f"scored {MIN_SCORE} or higher on a 0–5 scale measuring strength and "
        f"specificity of policy recommendation. Editorial review recommended "
        f"before publishing."
    )
    r.italic = True

    doc.add_paragraph(
        "Each piece is listed under its primary domain. Click any title to read "
        "the original."
    )

    for domain in DOMAIN_ORDER:
        items = by_domain.get(domain, [])
        if not items:
            continue
        items.sort(key=lambda q: -int(q["date"].replace("-", "")))
        doc.add_heading(f"{domain} ({len(items)})", level=1)

        if domain == "Public Safety":
            # Sub-group by what / who the ask is aimed at.
            sub = defaultdict(list)
            sub_order = [label for label, _ in PUBLIC_SAFETY_SUBGROUPS] + [
                "Other (federal, regulatory, etc.)"
            ]
            for q in items:
                sub[classify_public_safety(q["target"], q["ask"])].append(q)
            for label in sub_order:
                if not sub.get(label):
                    continue
                doc.add_heading(f"{label} ({len(sub[label])})", level=2)
                for q in sub[label]:
                    add_ask_paragraph(doc, q)
        else:
            for q in items:
                add_ask_paragraph(doc, q)

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUT))
    print(f"wrote {OUT.relative_to(ROOT)} ({total} articles across "
          f"{sum(1 for d in DOMAIN_ORDER if by_domain.get(d))} domains)")
    return 0


if __name__ == "__main__":
    sys.exit(main())
